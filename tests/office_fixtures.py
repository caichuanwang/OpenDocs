from __future__ import annotations

import io
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx.document import Document as WordDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.styles.style import ParagraphStyle
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from PIL import Image


def png_bytes(*, size: tuple[int, int] = (24, 24), color: str = "red") -> bytes:
    buffer = io.BytesIO()
    with Image.new("RGB", size, color) as image:
        image.save(buffer, format="PNG")
    return buffer.getvalue()


def save_document(document: WordDocument, path: Path) -> Path:
    document.save(str(path))
    return path


def add_hyperlink(
    paragraph: Paragraph,
    text: str,
    *,
    address: str | None = None,
    anchor: str | None = None,
) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    if address:
        relationship_id = paragraph.part.relate_to(
            address,
            RELATIONSHIP_TYPE.HYPERLINK,
            is_external=True,
        )
        hyperlink.set(qn("r:id"), relationship_id)
    if anchor:
        hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_picture(paragraph: Paragraph, image: bytes, *, alt_text: str | None = None) -> Run:
    run = paragraph.add_run()
    inline = run.add_picture(io.BytesIO(image), width=Inches(0.5))._inline
    if alt_text:
        inline.docPr.set("descr", alt_text)
    return run


def clone_run(paragraph: Paragraph, run: Run) -> None:
    paragraph._p.append(deepcopy(run._r))


def set_direct_outline_level(paragraph: Paragraph, level: int) -> None:
    outline = paragraph._p.get_or_add_pPr().get_or_add_outlineLvl()
    outline.set(qn("w:val"), str(max(level - 1, 0)))


def add_paragraph_style(
    document: WordDocument,
    name: str,
    *,
    base_style: str | ParagraphStyle | None = None,
) -> ParagraphStyle:
    if name in document.styles:
        style = document.styles[name]
    else:
        style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if base_style is not None:
        style.base_style = (
            document.styles[base_style] if isinstance(base_style, str) else base_style
        )
    return style


def add_numbering_definition(
    document: WordDocument,
    formats: list[str],
    *,
    starts: list[int] | None = None,
    overrides: dict[int, int] | None = None,
) -> int:
    numbering = document.part.numbering_part.numbering_definitions._numbering
    abstract_id = _next_numeric_id(numbering, qn("w:abstractNum"), qn("w:abstractNumId"))
    num_id = _next_numeric_id(numbering, qn("w:num"), qn("w:numId"))

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "multilevel")
    abstract.append(multi_level)
    for level, num_fmt in enumerate(formats):
        start = starts[level] if starts is not None else 1
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start_node = OxmlElement("w:start")
        start_node.set(qn("w:val"), str(start))
        fmt_node = OxmlElement("w:numFmt")
        fmt_node.set(qn("w:val"), num_fmt)
        text_node = OxmlElement("w:lvlText")
        text_node.set(qn("w:val"), "•" if num_fmt == "bullet" else f"%{level + 1}.")
        lvl.append(start_node)
        lvl.append(fmt_node)
        lvl.append(text_node)
        abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    for level, start_override in sorted((overrides or {}).items()):
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), str(level))
        start_node = OxmlElement("w:startOverride")
        start_node.set(qn("w:val"), str(start_override))
        override.append(start_node)
        num.append(override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph: Paragraph, *, num_id: int, level: int) -> None:
    num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    num_pr.get_or_add_numId().set(qn("w:val"), str(num_id))
    num_pr.get_or_add_ilvl().set(qn("w:val"), str(level))


def mark_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:tblHeader"))


def _next_numeric_id(parent: Any, tag: str, attribute: str) -> int:
    values = [
        int(child.get(attribute))
        for child in parent
        if child.tag == tag and child.get(attribute) and child.get(attribute).isdigit()
    ]
    return max(values, default=-1) + 1
