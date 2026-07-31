from __future__ import annotations

import hashlib
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_BREAK

from opendocs._models import (
    HeadingBlock,
    InlineLink,
    InlineText,
    ListItemBlock,
    ListKind,
    ParagraphBlock,
    SpannedTableBlock,
    TableBlock,
)
from opendocs.parsers.office.docx import extract_docx
from opendocs.parsers.office.models import BreakSlot, ImageSlot, NativeSlot, OfficePage
from opendocs.source import ParseWorkspace
from tests.office_fixtures import (
    add_hyperlink,
    add_numbering_definition,
    add_paragraph_style,
    add_picture,
    apply_numbering,
    clone_run,
    mark_table_header,
    png_bytes,
    save_document,
    set_direct_outline_level,
)


def _native_blocks(page: OfficePage) -> list[object]:
    return [block for slot in page.slots if isinstance(slot, NativeSlot) for block in slot.blocks]


def _paragraph_text(block: object) -> str:
    if not isinstance(block, ParagraphBlock | HeadingBlock | ListItemBlock):
        return ""
    return "".join(
        inline.text if isinstance(inline, InlineText) else inline.label for inline in block.inlines
    )


def _native_slot(slot: object) -> NativeSlot:
    assert isinstance(slot, NativeSlot)
    return slot


def test_extract_docx_preserves_body_order_inline_order_and_reused_media(tmp_path: Path) -> None:
    image = png_bytes()
    workspace = ParseWorkspace(tmp_path / "workspace")
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("intro ")
    add_hyperlink(paragraph, "safe", address="https://example.test")
    picture_run = add_picture(paragraph, image, alt_text="diagram")
    paragraph.add_run(" middle ")
    clone_run(paragraph, picture_run)
    paragraph.add_run("tail")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    document.add_paragraph("after")
    path = save_document(document, tmp_path / "ordered.docx")

    result = extract_docx(path, workspace)

    assert len(result.pages) == 1
    slots = result.pages[0].slots
    assert [type(slot) for slot in slots] == [
        NativeSlot,
        ImageSlot,
        NativeSlot,
        ImageSlot,
        NativeSlot,
        NativeSlot,
        NativeSlot,
    ]
    first, second, third = (
        _native_slot(slots[0]).blocks[0],
        _native_slot(slots[2]).blocks[0],
        _native_slot(slots[4]).blocks[0],
    )
    assert isinstance(first, ParagraphBlock)
    assert isinstance(second, ParagraphBlock)
    assert isinstance(third, ParagraphBlock)
    assert _paragraph_text(first) == "intro safe"
    assert [inline.target for inline in first.inlines if isinstance(inline, InlineLink)] == [
        "https://example.test"
    ]
    assert _paragraph_text(second) == " middle "
    assert _paragraph_text(third) == "tail"
    assert _native_slot(slots[5]).blocks == (TableBlock((("A", "B"), ("1", "2")), 0),)
    assert _paragraph_text(_native_slot(slots[6]).blocks[0]) == "after"

    image_slots = [slot for slot in slots if isinstance(slot, ImageSlot)]
    assert len(image_slots) == 2
    assert image_slots[0].artifact_name == image_slots[1].artifact_name
    assert image_slots[0].content_sha256 == hashlib.sha256(image).hexdigest()
    assert image_slots[0].content_sha256 == image_slots[1].content_sha256
    assert image_slots[0].alt_text == "diagram"
    assert workspace.output_path(image_slots[0].artifact_name).read_bytes() == image
    assert image_slots[0].bbox.require_normalized() == image_slots[0].bbox


def test_extract_docx_detects_direct_inherited_and_named_headings(tmp_path: Path) -> None:
    document = Document()
    direct = document.add_paragraph("Direct heading")
    set_direct_outline_level(direct, 1)
    inherited_style = add_paragraph_style(document, "Inherited Heading", base_style="Heading 2")
    inherited = document.add_paragraph("Inherited heading")
    inherited.style = inherited_style
    english_style = add_paragraph_style(document, "Heading 4")
    english = document.add_paragraph("English heading")
    english.style = english_style
    chinese_style = add_paragraph_style(document, "标题 3")
    chinese = document.add_paragraph("中文标题")
    chinese.style = chinese_style
    document.add_paragraph("Body")
    path = save_document(document, tmp_path / "headings.docx")

    result = extract_docx(path, ParseWorkspace(tmp_path / "workspace"))

    assert [
        (
            type(block),
            block.level if isinstance(block, HeadingBlock) else None,
            _paragraph_text(block),
        )
        for block in _native_blocks(result.pages[0])
    ] == [
        (HeadingBlock, 1, "Direct heading"),
        (HeadingBlock, 2, "Inherited heading"),
        (HeadingBlock, 4, "English heading"),
        (HeadingBlock, 3, "中文标题"),
        (ParagraphBlock, None, "Body"),
    ]


def test_extract_docx_resolves_nested_lists_restart_and_start_override(tmp_path: Path) -> None:
    document = Document()
    primary_num = add_numbering_definition(document, ["decimal", "bullet"])
    restart_num = add_numbering_definition(document, ["decimal"], overrides={0: 1})
    offset_num = add_numbering_definition(document, ["decimal"], overrides={0: 5})
    first = document.add_paragraph("One")
    apply_numbering(first, num_id=primary_num, level=0)
    nested = document.add_paragraph("Nested")
    apply_numbering(nested, num_id=primary_num, level=1)
    document.add_paragraph("Pause")
    second = document.add_paragraph("Two")
    apply_numbering(second, num_id=primary_num, level=0)
    restart = document.add_paragraph("Restart")
    apply_numbering(restart, num_id=restart_num, level=0)
    offset = document.add_paragraph("Five")
    apply_numbering(offset, num_id=offset_num, level=0)
    path = save_document(document, tmp_path / "lists.docx")

    result = extract_docx(path, ParseWorkspace(tmp_path / "workspace"))
    blocks = _native_blocks(result.pages[0])

    assert isinstance(blocks[0], ListItemBlock)
    assert (blocks[0].list_id, blocks[0].level, blocks[0].kind, blocks[0].ordinal) == (
        0,
        0,
        ListKind.ORDERED,
        1,
    )
    assert isinstance(blocks[1], ListItemBlock)
    assert (blocks[1].list_id, blocks[1].level, blocks[1].kind, blocks[1].ordinal) == (
        0,
        1,
        ListKind.BULLET,
        1,
    )
    assert isinstance(blocks[2], ParagraphBlock)
    assert _paragraph_text(blocks[2]) == "Pause"
    assert isinstance(blocks[3], ListItemBlock)
    assert (blocks[3].list_id, blocks[3].level, blocks[3].kind, blocks[3].ordinal) == (
        0,
        0,
        ListKind.ORDERED,
        2,
    )
    assert isinstance(blocks[4], ListItemBlock)
    assert (blocks[4].list_id, blocks[4].ordinal) == (1, 1)
    assert isinstance(blocks[5], ListItemBlock)
    assert (blocks[5].list_id, blocks[5].ordinal) == (2, 5)


def test_extract_docx_preserves_safe_unsafe_and_internal_links(tmp_path: Path) -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Links: ")
    add_hyperlink(paragraph, "safe", address="https://example.test")
    paragraph.add_run(" / ")
    add_hyperlink(paragraph, "unsafe", address="javascript:alert(1)")
    paragraph.add_run(" / ")
    add_hyperlink(paragraph, "jump", anchor="bookmark")
    path = save_document(document, tmp_path / "links.docx")

    result = extract_docx(path, ParseWorkspace(tmp_path / "workspace"))
    block = _native_blocks(result.pages[0])[0]

    assert isinstance(block, ParagraphBlock)
    assert _paragraph_text(block) == "Links: safe / unsafe / jump"
    assert [inline.target for inline in block.inlines if isinstance(inline, InlineLink)] == [
        "https://example.test",
        "javascript:alert(1)",
        "#bookmark",
    ]


def test_extract_docx_emits_tables_headers_merges_and_nested_flattening(tmp_path: Path) -> None:
    document = Document()
    simple = document.add_table(rows=2, cols=2)
    mark_table_header(simple.rows[0])
    simple.cell(0, 0).text = "H1"
    simple.cell(0, 1).text = "H2"
    simple.cell(1, 0).text = "A"
    simple.cell(1, 1).text = "B"

    merged = document.add_table(rows=3, cols=2)
    merged.cell(0, 0).merge(merged.cell(0, 1)).text = "wide"
    merged.cell(1, 0).merge(merged.cell(2, 0)).text = "tall"
    merged.cell(1, 1).text = "right-1"
    merged.cell(2, 1).text = "right-2"

    nested = document.add_table(rows=1, cols=2)
    nested.cell(0, 0).text = "outer"
    cell = nested.cell(0, 1)
    cell.text = "prefix"
    inner = cell.add_table(rows=1, cols=2)
    inner.cell(0, 0).text = "nested a"
    inner.cell(0, 1).text = "nested b"
    cell.add_paragraph("suffix")
    path = save_document(document, tmp_path / "tables.docx")

    result = extract_docx(path, ParseWorkspace(tmp_path / "workspace"))
    blocks = _native_blocks(result.pages[0])

    assert blocks[0] == TableBlock((("H1", "H2"), ("A", "B")), 1)
    assert isinstance(blocks[1], SpannedTableBlock)
    assert (blocks[1].row_count, blocks[1].column_count, blocks[1].header_rows) == (3, 2, 0)
    assert [(cell.text, cell.row_span, cell.column_span) for cell in blocks[1].cells] == [
        ("wide", 1, 2),
        ("tall", 2, 1),
        ("right-1", 1, 1),
        ("right-2", 1, 1),
    ]
    assert isinstance(blocks[2], TableBlock)
    assert blocks[2].grid[0][0] == "outer"
    assert "prefix" in blocks[2].grid[0][1]
    assert "nested a" in blocks[2].grid[0][1]
    assert "nested b" in blocks[2].grid[0][1]
    assert "suffix" in blocks[2].grid[0][1]
    assert [warning.code for warning in result.warnings] == ["docx_nested_table_flattened"]


def test_extract_docx_emits_manual_page_breaks_and_page_start_sections(tmp_path: Path) -> None:
    document = Document()
    paragraph = document.add_paragraph("before")
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    paragraph.add_run("after")
    document.add_section(WD_SECTION_START.NEW_PAGE)
    document.add_paragraph("after section")
    path = save_document(document, tmp_path / "breaks.docx")

    result = extract_docx(path, ParseWorkspace(tmp_path / "workspace"))
    slots = result.pages[0].slots

    assert [type(slot) for slot in slots] == [
        NativeSlot,
        BreakSlot,
        NativeSlot,
        BreakSlot,
        NativeSlot,
    ]
    assert _paragraph_text(_native_slot(slots[0]).blocks[0]) == "before"
    assert _paragraph_text(_native_slot(slots[2]).blocks[0]) == "after"
    assert _paragraph_text(_native_slot(slots[4]).blocks[0]) == "after section"


def test_extract_docx_empty_document_has_no_semantic_slots(tmp_path: Path) -> None:
    path = save_document(Document(), tmp_path / "empty.docx")

    result = extract_docx(path, ParseWorkspace(tmp_path / "workspace"))

    assert len(result.pages) == 1
    assert result.pages[0].slots == ()
    assert result.warnings == ()
